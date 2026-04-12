from __future__ import annotations

from pathlib import Path

from .errors import DocumentReadError, MissingDependencyError
from .models import CorpusOverview, DocumentMetadata, SourceDocument
from .text_utils import count_words, normalize_text

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def load_documents(input_dir: Path) -> list[SourceDocument]:
    if not input_dir.exists():
        input_dir.mkdir(parents=True, exist_ok=True)

    paths = sorted(
        path
        for path in input_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    documents: list[SourceDocument] = []
    for path in paths:
        text = read_document(path)
        normalized = normalize_text(text)
        documents.append(
            SourceDocument(
                path=str(path),
                extension=path.suffix.lower(),
                text=normalized,
                char_count=len(normalized),
                word_count=count_words(normalized),
            )
        )

    return documents


def read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _read_text_file(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise DocumentReadError(f"Unsupported document type: {path}")


def build_corpus_overview(documents: list[SourceDocument]) -> CorpusOverview:
    metadata = [
        DocumentMetadata(
            path=document.path,
            extension=document.extension,
            char_count=document.char_count,
            word_count=document.word_count,
        )
        for document in documents
    ]
    return CorpusOverview(
        source_count=len(documents),
        total_characters=sum(document.char_count for document in documents),
        total_words=sum(document.word_count for document in documents),
        sources=metadata,
    )


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise DocumentReadError(f"Unable to read text file: {path}") from exc


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingDependencyError(
            "Reading PDF files requires the 'pypdf' package."
        ) from exc

    try:
        reader = PdfReader(str(path))
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = normalize_text(page.extract_text() or "")
            if page_text:
                pages.append(f"[Page {index}]\n{page_text}")
        return "\n\n".join(pages)
    except Exception as exc:  # noqa: BLE001
        raise DocumentReadError(f"Unable to read PDF file: {path}") from exc


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise MissingDependencyError(
            "Reading DOCX files requires the 'python-docx' package."
        ) from exc

    try:
        document = Document(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as exc:  # noqa: BLE001
        raise DocumentReadError(f"Unable to read DOCX file: {path}") from exc
