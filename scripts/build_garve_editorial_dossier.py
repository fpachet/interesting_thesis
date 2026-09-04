#!/usr/bin/env python3
"""Build a polished DOCX from the Garve editorial dossier Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/dossier-editorial-garve.md"
OUTPUT = ROOT / "docs/Dossier_editorial_Garve.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 36, 48)
GRAY = RGBColor(92, 101, 110)
LIGHT_GRAY = "F4F6F9"


def set_font(run, *, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_spacing(paragraph, *, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_bottom_border(paragraph, color="D9E2EC", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    style_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Lead" not in [s.name for s in doc.styles]:
        lead = doc.styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = doc.styles["Lead"]
    lead.base_style = normal
    lead.font.name = "Calibri"
    lead._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    lead._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    lead.font.size = Pt(11.5)
    lead.font.color.rgb = DARK_BLUE
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(12)
    lead.paragraph_format.line_spacing = 1.25

    configure_numbering(doc)
    configure_header_footer(section)


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, after=3, line=1.0)
    left = p.add_run("PROJET D’ÉDITION  |  CHRISTIAN GARVE")
    set_font(left, size=8.5, bold=True, color=GRAY)
    set_bottom_border(p)

    footer = section.footer
    p = footer.paragraphs[0]
    set_spacing(p, before=3, line=1.0)
    add_page_field(p)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    next_abstract = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def add_definition(abstract_id, num_id, fmt, text, align_twips, indent_twips, hanging_twips):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(indent_twips))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_twips))
        ind.set(qn("w:hanging"), str(hanging_twips))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        numbering.append(num)

    add_definition(next_abstract, next_num, "bullet", "•", 260, 540, 280)
    add_definition(next_abstract + 1, next_num + 1, "decimal", "%1.", 260, 540, 280)
    doc._garve_bullet_num_id = next_num
    doc._garve_decimal_num_id = next_num + 1
    doc._garve_decimal_abstract_id = next_abstract + 1


def new_decimal_num_id(doc):
    """Create a fresh numbering instance so each independent list restarts at 1."""
    numbering = doc.part.numbering_part.element
    existing = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(doc._garve_decimal_abstract_id))
    num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_inline(paragraph, text):
    """Add a small, deterministic subset of Markdown inline formatting."""
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, bold=True, color=INK)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, italic=True, color=INK)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, color=INK)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=28, after=8, line=1.0)
    run = p.add_run("CHRISTIAN GARVE")
    set_font(run, size=11, bold=True, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=8, line=1.0)
    run = p.add_run("Quelques pensées sur l’intéressant")
    set_font(run, size=26, bold=True, italic=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=20, line=1.0)
    run = p.add_run("Projet de traduction française, avec introduction et annotations")
    set_font(run, size=14, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=4, line=1.0)
    run = p.add_run("François Pachet")
    set_font(run, size=11.5, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=22, line=1.0)
    run = p.add_run("Dossier de travail — 26 août 2026")
    set_font(run, size=10, italic=True, color=GRAY)

    p = doc.add_paragraph()
    set_spacing(p, after=12, line=1.15)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    label = p.add_run("TEXTE SOURCE  ")
    set_font(label, size=9.5, bold=True, color=DARK_BLUE)
    text = (
        "Christian Garve, « Einige Gedanken über das Interessirende », publié en 1771-1772, "
        "puis repris avec une annexe dans la Sammlung einiger Abhandlungen, Leipzig, 1779, p. 253-439."
    )
    run = p.add_run(text)
    set_font(run, size=9.5, color=INK)


def parse_body(doc, lines):
    paragraph_buffer = []
    first_body_para = True
    current_decimal_num_id = None

    def flush_paragraph():
        nonlocal paragraph_buffer, first_body_para
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        p = doc.add_paragraph(style="Lead" if first_body_para else "Normal")
        add_inline(p, text)
        first_body_para = False
        paragraph_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line:
            flush_paragraph()
            current_decimal_num_id = None
            continue
        if line.startswith("### "):
            flush_paragraph()
            current_decimal_num_id = None
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:])
        elif line.startswith("## "):
            flush_paragraph()
            current_decimal_num_id = None
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:])
        elif re.match(r"^- ", line):
            flush_paragraph()
            current_decimal_num_id = None
            p = doc.add_paragraph(style="Normal")
            apply_num(p, doc._garve_bullet_num_id)
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            flush_paragraph()
            if current_decimal_num_id is None:
                current_decimal_num_id = new_decimal_num_id(doc)
            p = doc.add_paragraph(style="Normal")
            apply_num(p, current_decimal_num_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        else:
            current_decimal_num_id = None
            paragraph_buffer.append(line)
    flush_paragraph()


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    # The first five non-empty lines are rendered through the custom title block.
    body_start = next(i for i, line in enumerate(lines) if line == "## Résumé du projet")

    doc = Document()
    configure_document(doc)
    core = doc.core_properties
    core.title = "Christian Garve — Projet de traduction française annotée"
    core.subject = "Dossier éditorial"
    core.author = "François Pachet"
    core.keywords = "Christian Garve, intéressant, traduction, philosophie, annotation"
    core.comments = "Dossier de travail préparé le 26 août 2026"

    add_cover(doc)
    parse_body(doc, lines[body_start:])
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
